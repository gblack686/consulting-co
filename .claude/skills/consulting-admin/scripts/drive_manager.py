"""
Google Drive operations for consulting admin.
Creates client folders, uploads documents, manages sharing.
"""
import os
import io
from googleapiclient.http import MediaFileUpload, MediaIoBaseUpload
from googleapiclient.errors import HttpError
from . import google_client


ROOT_FOLDER_NAME = "GBAutomation Clients"
ONBOARDING_FOLDER_NAME = "Onboarding"


def get_or_create_folder(name: str, parent_id: str = None) -> str:
    """Get an existing folder by name (under parent) or create it. Returns folder ID."""
    drive = google_client.drive_service()

    query = f"name='{name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
    if parent_id:
        query += f" and '{parent_id}' in parents"

    results = drive.files().list(q=query, fields="files(id, name)").execute()
    files = results.get("files", [])

    if files:
        return files[0]["id"]

    metadata = {
        "name": name,
        "mimeType": "application/vnd.google-apps.folder",
    }
    if parent_id:
        metadata["parents"] = [parent_id]

    folder = drive.files().create(body=metadata, fields="id").execute()
    return folder["id"]


def get_root_folder_id() -> str:
    return get_or_create_folder(ROOT_FOLDER_NAME)


def create_client_folder(client_name: str) -> str:
    """Create GBAutomation Clients / {client_name} folder. Returns folder ID."""
    root_id = get_root_folder_id()
    return get_or_create_folder(client_name, parent_id=root_id)


def create_onboarding_folder(client_name: str) -> tuple[str, str]:
    """
    Create GBAutomation Clients / {client_name} / Onboarding folder.
    Returns (client_folder_id, onboarding_folder_id).
    """
    client_folder_id = create_client_folder(client_name)
    onboarding_folder_id = get_or_create_folder(ONBOARDING_FOLDER_NAME, parent_id=client_folder_id)
    return client_folder_id, onboarding_folder_id


def get_files_shared_by(owner_email: str) -> list[dict]:
    """
    Find all Drive files shared with greg@gbautomation.xyz that are owned by owner_email.
    Returns list of dicts: {id, name, mimeType, webViewLink}.
    """
    drive = google_client.drive_service()
    results = []
    page_token = None

    while True:
        resp = drive.files().list(
            q="sharedWithMe=true",
            fields="nextPageToken, files(id, name, mimeType, webViewLink, owners, sharingUser)",
            pageSize=100,
            pageToken=page_token,
        ).execute()

        for f in resp.get("files", []):
            owners = [o.get("emailAddress", "").lower() for o in f.get("owners", [])]
            sharer = (f.get("sharingUser") or {}).get("emailAddress", "").lower()
            if owner_email.lower() in owners or owner_email.lower() == sharer:
                results.append({
                    "id": f["id"],
                    "name": f["name"],
                    "mimeType": f["mimeType"],
                    "webViewLink": f.get("webViewLink", ""),
                })

        page_token = resp.get("nextPageToken")
        if not page_token:
            break

    return results


def export_doc_text(file_id: str, mime_type: str) -> str:
    """
    Export readable text from a Drive file.
    Supports Google Docs, Sheets, Slides, and plain text files.
    Returns extracted text or a note if unsupported.
    """
    drive = google_client.drive_service()

    export_map = {
        "application/vnd.google-apps.document": "text/plain",
        "application/vnd.google-apps.spreadsheet": "text/csv",
        "application/vnd.google-apps.presentation": "text/plain",
    }

    if mime_type in export_map:
        content = drive.files().export(
            fileId=file_id,
            mimeType=export_map[mime_type]
        ).execute()
        return content.decode("utf-8", errors="replace") if isinstance(content, bytes) else str(content)

    elif mime_type == "text/plain":
        content = drive.files().get_media(fileId=file_id).execute()
        return content.decode("utf-8", errors="replace") if isinstance(content, bytes) else str(content)

    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx_text(drive, file_id)

    elif mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        return _extract_xlsx_text(drive, file_id)

    elif mime_type == "application/msword":
        return _extract_docx_text(drive, file_id)

    elif mime_type == "application/pdf":
        return "[PDF — manual review needed, cannot auto-extract text]"

    elif mime_type == "application/vnd.google-apps.shortcut":
        return "[Shortcut — follow link manually to access target file]"

    else:
        return f"[{mime_type} — unsupported format, skipped]"


def _extract_docx_text(drive, file_id: str) -> str:
    """Download a .docx file from Drive and extract plain text."""
    try:
        import docx
        raw = drive.files().get_media(fileId=file_id).execute()
        doc = docx.Document(io.BytesIO(raw))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table cells
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[docx extraction failed: {e}]"


def _extract_xlsx_text(drive, file_id: str) -> str:
    """Download a .xlsx file from Drive and extract text from all sheets."""
    try:
        import openpyxl
        raw = drive.files().get_media(fileId=file_id).execute()
        wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        rows = []
        for sheet in wb.worksheets:
            rows.append(f"=== Sheet: {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    rows.append(" | ".join(cells))
        return "\n".join(rows)
    except Exception as e:
        return f"[xlsx extraction failed: {e}]"


def copy_file_to_folder(file_id: str, new_name: str, parent_folder_id: str) -> str:
    """Copy a Drive file into a target folder. Returns new file's webViewLink."""
    drive = google_client.drive_service()
    copy = drive.files().copy(
        fileId=file_id,
        body={"name": new_name, "parents": [parent_folder_id]},
        fields="id, webViewLink"
    ).execute()
    return copy["webViewLink"]


def create_google_doc(title: str, content_markdown: str, parent_folder_id: str, logo_path: str = None) -> str:
    """
    Create a Google Doc with the given title and body content.
    Optionally inserts a logo at the top.
    Returns the document URL.
    """
    drive = google_client.drive_service()
    docs = google_client.docs_service()

    # Create empty doc
    doc_metadata = {
        "name": title,
        "mimeType": "application/vnd.google-apps.document",
        "parents": [parent_folder_id],
    }
    doc = drive.files().create(body=doc_metadata, fields="id, webViewLink").execute()
    doc_id = doc["id"]

    # Build requests to populate the doc
    requests = []

    # Insert main content text (at end of doc)
    requests.append({
        "insertText": {
            "location": {"index": 1},
            "text": content_markdown
        }
    })

    if requests:
        docs.documents().batchUpdate(
            documentId=doc_id,
            body={"requests": requests}
        ).execute()

    # If logo provided, insert it at the top
    if logo_path and os.path.exists(logo_path):
        _insert_logo(docs, drive, doc_id, logo_path)

    return doc["webViewLink"]


def _insert_logo(docs, drive, doc_id: str, logo_path: str):
    """Upload logo to Drive (temp) and insert inline at top of doc."""
    # Upload logo as a temp file
    media = MediaFileUpload(logo_path, mimetype="image/png")
    logo_file = drive.files().create(
        body={"name": "gb-logo-temp.png"},
        media_body=media,
        fields="id, webContentLink"
    ).execute()

    # Make it publicly readable so Docs API can fetch it
    drive.permissions().create(
        fileId=logo_file["id"],
        body={"type": "anyone", "role": "reader"}
    ).execute()

    img_url = f"https://drive.google.com/uc?id={logo_file['id']}"

    # Insert image at index 1 (top of document)
    docs.documents().batchUpdate(
        documentId=doc_id,
        body={"requests": [
            {
                "insertInlineImage": {
                    "location": {"index": 1},
                    "uri": img_url,
                    "objectSize": {
                        "height": {"magnitude": 60, "unit": "PT"},
                        "width": {"magnitude": 200, "unit": "PT"},
                    }
                }
            }
        ]}
    ).execute()

    # Clean up temp logo file
    drive.files().delete(fileId=logo_file["id"]).execute()


def share_folder(folder_id: str, email: str, role: str = "reader"):
    """Share a Drive folder with an external email."""
    drive = google_client.drive_service()
    try:
        drive.permissions().create(
            fileId=folder_id,
            body={
                "type": "user",
                "role": role,
                "emailAddress": email,
            },
            sendNotificationEmail=False,
            fields="id"
        ).execute()
    except Exception:
        # Non-Google accounts require sendNotificationEmail=True
        drive.permissions().create(
            fileId=folder_id,
            body={
                "type": "user",
                "role": role,
                "emailAddress": email,
            },
            sendNotificationEmail=True,
            fields="id"
        ).execute()


def get_folder_url(folder_id: str) -> str:
    return f"https://drive.google.com/drive/folders/{folder_id}"


def upload_file_to_folder(local_path: str, filename: str, folder_id: str, mime_type: str = "application/vnd.openxmlformats-officedocument.presentationml.presentation") -> str:
    """Upload a local file to a Drive folder. Returns the file URL."""
    drive = google_client.drive_service()
    file_metadata = {"name": filename, "parents": [folder_id]}
    media = MediaFileUpload(local_path, mimetype=mime_type, resumable=False)
    f = drive.files().create(body=file_metadata, media_body=media, fields="id").execute()
    return f"https://drive.google.com/file/d/{f['id']}/view"
